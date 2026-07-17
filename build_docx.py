# -*- coding: utf-8 -*-
"""
build_docx.py - Conversor Markdown -> DOCX para o livro "English Through My Routine".

Uso:
    python build_docx.py src/00-Plano-Editorial.md livro/00-Plano-Editorial.docx

Suporta: # ## ### #### headings, **negrito**, *itálico*, `código`,
listas (- / * / 1.), tabelas |a|b|, > citação, --- linha, ``` blocos de código.
Feito para gerar material didático bonito e pronto para impressão.
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta ----------
NAVY   = RGBColor(0x1F, 0x33, 0x55)   # títulos
BLUE   = RGBColor(0x2E, 0x5E, 0xAA)   # subtítulos
ACCENT = RGBColor(0xC0, 0x39, 0x2B)   # destaques
GREY   = RGBColor(0x55, 0x55, 0x55)
BODY_FONT = "Calibri"
HEAD_FONT = "Calibri"


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E5EAA')
    pbdr.append(bottom)
    pPr.append(pbdr)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')


def add_inline(paragraph, text, base_color=None, base_size=None, base_bold=False):
    """Escreve texto com **negrito**, *itálico* e `código` em runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        if part.startswith('**') and part.endswith('**'):
            run.text = part[2:-2]; run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run.text = part[1:-1]; run.font.name = 'Consolas'
            run.font.color.rgb = ACCENT
        elif part.startswith('*') and part.endswith('*'):
            run.text = part[1:-1]; run.italic = True
        else:
            run.text = part
        if base_bold:
            run.bold = True
        if base_color is not None:
            run.font.color.rgb = base_color
        if base_size is not None:
            run.font.size = Pt(base_size)
        if run.font.name is None:
            run.font.name = BODY_FONT


def style_base(doc):
    st = doc.styles['Normal']
    st.font.name = BODY_FONT
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)


def build_table(doc, rows):
    # rows: lista de listas de células (strings). Primeira linha = cabeçalho.
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            txt = row[j] if j < len(row) else ''
            cell = cells[j]
            cell.paragraphs[0].text = ''
            add_inline(cell.paragraphs[0], txt, base_bold=(i == 0))
            if i == 0:
                set_cell_bg(cell, '1F3355')
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    style_base(doc)

    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]

        # blocos de código ```
        if line.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buf))
                run.font.name = 'Consolas'; run.font.size = Pt(10)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        stripped = line.strip()

        # linha horizontal
        if stripped in ('---', '***', '___'):
            add_hr(doc); i += 1; continue

        # tabela
        if stripped.startswith('|') and '|' in stripped[1:]:
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i].strip()); i += 1
            rows = []
            for b in block:
                if re.match(r'^\|[\s:|-]+\|?$', b):  # linha separadora ---|---
                    continue
                cells = [c.strip() for c in b.strip('|').split('|')]
                rows.append(cells)
            if rows:
                build_table(doc, rows)
            continue

        # headings
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1)); text = m.group(2).strip()
            if level == 1:
                h = doc.add_heading('', level=0)
                add_inline(h, text, base_color=NAVY, base_size=24, base_bold=True)
            elif level == 2:
                h = doc.add_heading('', level=1)
                add_inline(h, text, base_color=BLUE, base_size=17, base_bold=True)
            elif level == 3:
                h = doc.add_heading('', level=2)
                add_inline(h, text, base_color=BLUE, base_size=13.5, base_bold=True)
            else:
                h = doc.add_heading('', level=3)
                add_inline(h, text, base_color=GREY, base_size=12, base_bold=True)
            i += 1; continue

        # citação
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, stripped[2:], base_color=GREY)
            for r in p.runs:
                r.italic = True
            i += 1; continue

        # lista com marcador
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, text)
            i += 1; continue

        # lista numerada
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            add_inline(p, text)
            i += 1; continue

        # linha em branco
        if stripped == '':
            i += 1; continue

        # parágrafo normal
        p = doc.add_paragraph()
        add_inline(p, line)
        i += 1

    import os
    os.makedirs(os.path.dirname(docx_path) or '.', exist_ok=True)
    doc.save(docx_path)
    print('OK ->', docx_path)


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
